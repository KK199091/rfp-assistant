import streamlit as st
import requests
import io
import os
import time
from PyPDF2 import PdfReader
import tempfile
from dotenv import load_dotenv
import json
import pandas as pd
import re

# Load environment variables from .env file
load_dotenv()

def generate_pdf(response_draft, review):
    """Generate a PDF document from the RFP response and review"""
    from reportlab.lib.pagesizes import letter
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from io import BytesIO
    
    # Create a BytesIO buffer to receive the PDF data
    buffer = BytesIO()
    
    # Create the PDF document
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    
    # Create custom styles
    title_style = ParagraphStyle(
        'Title',
        parent=styles['Heading1'],
        fontSize=16,
        textColor=colors.HexColor('#1e3a8a'),
        spaceAfter=12
    )
    
    heading2_style = ParagraphStyle(
        'Heading2',
        parent=styles['Heading2'],
        fontSize=14,
        textColor=colors.HexColor('#1e40af'),
        spaceAfter=10
    )
    
    normal_style = ParagraphStyle(
        'Normal',
        parent=styles['Normal'],
        fontSize=11,
        spaceAfter=8
    )
    
    # Build the document content
    content = []
    
    # Add title
    content.append(Paragraph("RFP Response Draft", title_style))
    content.append(Spacer(1, 12))
    
    # Basic markdown to PDF conversion (simplified)
    lines = response_draft.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('# '):
            content.append(Paragraph(line[2:], title_style))
        elif line.startswith('## '):
            content.append(Paragraph(line[3:], heading2_style))
        elif line.startswith('### '):
            content.append(Paragraph(line[4:], heading2_style))
        elif line.startswith('- ') or line.startswith('* '):
            content.append(Paragraph(f"• {line[2:]}", normal_style))
        else:
            content.append(Paragraph(line, normal_style))
    
    # Add review section
    content.append(Spacer(1, 20))
    content.append(Paragraph("Quality Review", title_style))
    content.append(Spacer(1, 12))
    
    # Add review content
    lines = review.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('# '):
            content.append(Paragraph(line[2:], title_style))
        elif line.startswith('## '):
            content.append(Paragraph(line[3:], heading2_style))
        elif line.startswith('### '):
            content.append(Paragraph(line[4:], heading2_style))
        elif line.startswith('- ') or line.startswith('* '):
            content.append(Paragraph(f"• {line[2:]}", normal_style))
        else:
            content.append(Paragraph(line, normal_style))
    
    # Build the PDF
    doc.build(content)
    buffer.seek(0)
    return buffer

def generate_docx(response_draft, review):
    """Generate a Word document from the RFP response and review"""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from io import BytesIO
    
    # Create a new Document
    doc = Document()
    
    # Add a title
    title = doc.add_heading("RFP Response Draft", level=1)
    title_format = title.runs[0].font
    title_format.size = Pt(18)
    
    # Process the response draft content
    lines = response_draft.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('# '):
            heading = doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            heading = doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            heading = doc.add_heading(line[4:], level=3)
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph()
            p.add_run("• " + line[2:])
            p.style = 'List Bullet'
        else:
            p = doc.add_paragraph(line)
    
    # Add a page break before the review section
    doc.add_page_break()
    
    # Add the review section title
    review_title = doc.add_heading("Quality Review", level=1)
    
    # Process the review content
    lines = review.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('# '):
            heading = doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            heading = doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            heading = doc.add_heading(line[4:], level=3)
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph()
            p.add_run("• " + line[2:])
            p.style = 'List Bullet'
        else:
            p = doc.add_paragraph(line)
    
    # Save the document to a BytesIO buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Password protection
def check_password():
    """Returns `True` if the user had the correct password."""

    def password_entered():
        """Checks whether a password entered by the user is correct."""
        if st.session_state["password"] == "rfpdriteam2025":
            st.session_state["password_correct"] = True
            del st.session_state["password"]  # Don't store the password
        else:
            st.session_state["password_correct"] = False

    # Return True if the password is validated
    if st.session_state.get("password_correct", False):
        return True

    # Show input for password
    st.markdown("""
    <style>
    .password-container {
        background-color: #f8f9fa;
        padding: 30px;
        border-radius: 10px;
        box-shadow: 0 4px 6px rgba(0,0,0,0.1);
        max-width: 500px;
        margin: 100px auto;
        text-align: center;
    }
    </style>
    <div class="password-container">
        <h1>RFP Response Assistant</h1>
        <h3>Restricted Access</h3>
        <p>This tool is only available to authorized team members.</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Create password input field
    st.text_input(
        "Enter the access password", 
        type="password", 
        on_change=password_entered, 
        key="password"
    )
    return False

# Now check if the password is correct
if not check_password():
    st.stop()  # Stop execution if password is incorrect

# Configure the Streamlit page
st.set_page_config(
    page_title="RFP Response Assistant",
    page_icon="📝",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Get API key from environment variable or Streamlit secrets
def get_api_key():
    # First try to get from Streamlit secrets (when deployed)
    if 'ANTHROPIC_API_KEY' in st.secrets:
        return st.secrets['ANTHROPIC_API_KEY']
    # Then try to get from environment variables (local development)
    elif os.getenv('ANTHROPIC_API_KEY'):
        return os.getenv('ANTHROPIC_API_KEY')
    else:
        st.error("API key not found. Please set it in .env file or Streamlit secrets.")
        st.stop()

# Initialize API key
api_key = get_api_key()

# Updated Anthropic API function for Claude 3 models using Messages API
def call_anthropic_api(prompt, max_tokens=2000, temperature=0, system_prompt=None):
    headers = {
        "x-api-key": api_key,
        "content-type": "application/json",
        "anthropic-version": "2023-06-01"
    }
    
    messages = [{"role": "user", "content": prompt}]
    
    data = {
        "model": "claude-3-haiku-20240307",  # Using Claude 3 Haiku which is fast and reliable
        "max_tokens": max_tokens,
        "temperature": temperature,
        "messages": messages
    }
    
    # Add system prompt if provided
    if system_prompt:
        data["system"] = system_prompt
    
    response = requests.post(
        "https://api.anthropic.com/v1/messages", 
        headers=headers,
        json=data
    )
    
    if response.status_code != 200:
        raise Exception(f"Error from Anthropic API: {response.text}")
    
    # Extract the response text from the messages API format
    response_json = response.json()
    return response_json["content"][0]["text"]

# Enhanced CSS with animations
st.markdown("""
<style>
    /* Modern UI Theme - Base */
    .main {
        background-color: #f8fafc;
        background-image: linear-gradient(135deg, #f8fafc 0%, #f1f5f9 100%);
        font-family: 'Segoe UI', -apple-system, BlinkMacSystemFont, sans-serif;
    }
    .stApp {
        max-width: 1200px;
        margin: 0 auto;
    }
    
    /* Header Styling */
    h1 {
        color: #1e3a8a;
        font-weight: 700;
        font-size: 2.5rem;
        margin-bottom: 0.5rem;
        padding-bottom: 0.5rem;
        border-bottom: 2px solid #4f46e5;
        display: inline-block;
    }
    h2 {
        color: #1e40af;
        font-weight: 600;
        margin-top: 1.5rem;
        margin-bottom: 1rem;
    }
    h3 {
        color: #1e3a8a;
        font-weight: 600;
        margin-top: 1.25rem;
    }
    
    /* Button Styling */
    .stButton>button {
        background-color: #4f46e5;
        color: white;
        padding: 0.75rem 1.5rem;
        border-radius: 8px;
        border: none;
        font-size: 1rem;
        font-weight: 600;
        transition: all 0.3s ease;
        box-shadow: 0 4px 6px rgba(79, 70, 229, 0.25);
    }
    .stButton>button:hover {
        background-color: #4338ca;
        box-shadow: 0 6px 10px rgba(79, 70, 229, 0.3);
        transform: translateY(-2px);
    }
    .stButton>button:active {
        transform: translateY(0);
    }
    
    /* Hero Section */
    .hero-container {
        display: flex;
        align-items: center;
        background: linear-gradient(135deg, #4f46e5 0%, #7c3aed 100%);
        border-radius: 16px;
        padding: 3rem 2rem;
        margin-bottom: 2rem;
        color: white;
        box-shadow: 0 20px 25px -5px rgba(0, 0, 0, 0.1), 0 10px 10px -5px rgba(0, 0, 0, 0.04);
        overflow: hidden;
        position: relative;
    }
    .hero-container::after {
        content: "";
        position: absolute;
        width: 300px;
        height: 300px;
        background: rgba(255, 255, 255, 0.1);
        border-radius: 50%;
        right: -100px;
        top: -100px;
    }
    .hero-content {
        flex: 1;
        z-index: 1;
    }
    .hero-content h1 {
        font-size: 3rem;
        font-weight: 800;
        margin-bottom: 1rem;
        color: white;
        border: none;
    }
    .hero-subtitle {
        font-size: 1.5rem;
        margin-bottom: 2rem;
        opacity: 0.9;
    }
    .hero-stats {
        display: flex;
        gap: 2rem;
        margin-top: 2rem;
    }
    .stat-item {
        display: flex;
        flex-direction: column;
    }
    .stat-number {
        font-size: 2rem;
        font-weight: 700;
    }
    .stat-label {
        font-size: 1rem;
        opacity: 0.8;
    }
    .hero-image {
        flex: 1;
        display: flex;
        justify-content: center;
        align-items: center;
    }
    
    /* Steps Section */
    .steps-container {
        display: flex;
        gap: 1.5rem;
        margin: 2rem 0;
    }
    .step-item {
        flex: 1;
        background-color: white;
        border-radius: 12px;
        padding: 2rem 1.5rem;
        text-align: center;
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.05);
        position: relative;
        transition: all 0.3s ease;
    }
    .step-item:hover {
        transform: translateY(-5px);
        box-shadow: 0 10px 15px rgba(0, 0, 0, 0.1);
    }
    .step-number {
        position: absolute;
        top: -15px;
        left: 50%;
        transform: translateX(-50%);
        background: #4f46e5;
        color: white;
        width: 30px;
        height: 30px;
        border-radius: 50%;
        display: flex;
        align-items: center;
        justify-content: center;
        font-weight: bold;
    }
    .step-icon {
        font-size: 2.5rem;
        margin-bottom: 1rem;
    }
    .step-title {
        font-weight: 600;
        font-size: 1.2rem;
        margin-bottom: 0.5rem;
        color: #1e40af;
    }
    .step-desc {
        color: #6b7280;
        font-size: 0.95rem;
    }
    
    /* Upload Container */
    .upload-container {
        background-color: white;
        border-radius: 16px;
        padding: 2rem;
        margin-top: 2rem;
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.05);
        border-top: 4px solid #4f46e5;
    }
    
    /* Info Boxes */
    .info-box {
        background-color: #e0f2fe;
        padding: 1.5rem;
        border-radius: 10px;
        margin: 1rem 0;
        border-left: 5px solid #0ea5e9;
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.05);
    }
    .success-box {
        background-color: #dcfce7;
        padding: 1.5rem;
        border-radius: 10px;
        margin: 1rem 0;
        border-left: 5px solid #10b981;
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.05);
    }
    .warning-box {
        background-color: #fef9c3;
        padding: 1.5rem;
        border-radius: 10px;
        margin: 1rem 0;
        border-left: 5px solid #eab308;
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.05);
    }
    
    /* Agent Cards with Animation */
    .agent-card {
        background-color: white;
        border-radius: 12px;
        padding: 1.5rem;
        margin-bottom: 1.5rem;
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.05);
        border-top: 4px solid #4f46e5;
        transition: all 0.3s ease;
        animation: slideIn 0.5s ease forwards;
        opacity: 0;
    }
    @keyframes slideIn {
        from {
            opacity: 0;
            transform: translateY(20px);
        }
        to {
            opacity: 1;
            transform: translateY(0);
        }
    }
    .agent-card:nth-child(1) { animation-delay: 0.1s; }
    .agent-card:nth-child(2) { animation-delay: 0.3s; }
    .agent-card:nth-child(3) { animation-delay: 0.5s; }
    .agent-card:nth-child(4) { animation-delay: 0.7s; }
    
    .agent-header {
        font-weight: 700;
        color: #1e40af;
        margin-bottom: 0.5rem;
        font-size: 1.25rem;
        display: flex;
        align-items: center;
    }
    .agent-status {
        font-style: italic;
        color: #6b7280;
        margin-bottom: 1rem;
    }
    
    /* Progress Animation */
    @keyframes progress {
        0% { width: 0%; }
        100% { width: 100%; }
    }
    .animated-progress {
        height: 4px;
        background: #e5e7eb;
        border-radius: 4px;
        margin: 1rem 0;
        overflow: hidden;
    }
    .animated-progress-bar {
        height: 100%;
        background: linear-gradient(90deg, #4f46e5, #7c3aed);
        animation: progress 30s ease forwards;
    }
    
    /* Success Animation */
    @keyframes checkmark {
        0% {
            stroke-dashoffset: 100;
        }
        100% {
            stroke-dashoffset: 0;
        }
    }
    .checkmark {
        width: 56px;
        height: 56px;
        border-radius: 50%;
        display: block;
        stroke-width: 2;
        stroke: #4f46e5;
        stroke-miterlimit: 10;
        margin: 10% auto;
        box-shadow: inset 0px 0px 0px #4f46e5;
    }
    .checkmark-circle {
        stroke-dasharray: 166;
        stroke-dashoffset: 166;
        stroke-width: 2;
        stroke-miterlimit: 10;
        stroke: #4f46e5;
        fill: none;
        animation: checkmark 0.6s cubic-bezier(0.65, 0, 0.45, 1) forwards;
    }
    .checkmark-check {
        transform-origin: 50% 50%;
        stroke-dasharray: 48;
        stroke-dashoffset: 48;
        animation: checkmark 0.3s cubic-bezier(0.65, 0, 0.45, 1) 0.3s forwards;
    }
    
    /* Tables and Data Display */
    .dataframe {
        border-collapse: collapse;
        width: 100%;
        border-radius: 8px;
        overflow: hidden;
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.05);
    }
    .dataframe th {
        background-color: #4f46e5;
        color: white;
        text-align: left;
        padding: 12px;
    }
    .dataframe td {
        padding: 10px 12px;
        border-bottom: 1px solid #e5e7eb;
    }
    .dataframe tr:nth-child(even) {
        background-color: #f9fafb;
    }
    .dataframe tr:hover {
        background-color: #f3f4f6;
    }
    
    /* Sidebar styling */
    .css-1d391kg {
        background-color: #f1f5f9;
    }
    .css-1544g2n {
        padding: 2rem 1rem;
    }
    
    /* Tabs styling */
    .stTabs [data-baseweb="tab-list"] {
        gap: 8px;
    }
    .stTabs [data-baseweb="tab"] {
        background-color: #f1f5f9;
        border-radius: 8px 8px 0 0;
        padding: 10px 20px;
        border: none;
    }
    .stTabs [aria-selected="true"] {
        background-color: #4f46e5 !important;
        color: white !important;
    }
    .stTabs [data-baseweb="tab"]:hover {
        transform: translateY(-2px);
    }
    
    /* Progress bar */
    .stProgress > div > div > div > div {
        background-color: #4f46e5;
    }
    
    /* File uploader */
    .stFileUploader > div > div {
        border: 2px dashed #4f46e5;
        border-radius: 10px;
        padding: 20px;
    }
    
    /* Download button */
    .stDownloadButton > button {
        background-color: #059669;
        transition: all 0.3s ease;
    }
    .stDownloadButton > button:hover {
        background-color: #047857;
        transform: translateY(-2px);
        box-shadow: 0 4px 6px rgba(4, 120, 87, 0.3);
    }
    
    /* Card-like sections */
    .card {
        background-color: white;
        border-radius: 12px;
        padding: 1.5rem;
        margin-bottom: 1.5rem;
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.05);
        animation: fadeIn 0.5s ease forwards;
        opacity: 0;
    }
    @keyframes fadeIn {
        from { opacity: 0; }
        to { opacity: 1; }
    }
    
    /* JSON formatting for requirements */
    .json-key {
        color: #0284c7;
        font-weight: 600;
    }
    .json-value {
        color: #4b5563;
    }
    .json-list-item {
        margin-left: 20px;
        padding: 6px 0;
        border-bottom: 1px dashed #e5e7eb;
    }
    .json-list-item:last-child {
        border-bottom: none;
    }
    
    /* Dashboard Styles */
    .dashboard-header {
        text-align: center;
        margin-bottom: 2rem;
    }
    .dashboard-header h2 {
        font-size: 1.8rem;
        color: #1e3a8a;
        margin-bottom: 0.5rem;
    }
    .dashboard-header p {
        color: #6b7280;
        font-size: 1.1rem;
    }
    .dashboard-card {
        background-color: white;
        border-radius: 12px;
        padding: 1.5rem;
        height: 100%;
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.05);
    }
    
    /* Metrics Bar */
    .metrics-container {
        display: flex;
        flex-wrap: wrap;
        gap: 1rem;
        margin-top: 1rem;
    }
    .metric-item {
        flex: 1;
        min-width: 120px;
        background-color: #f8fafc;
        border-radius: 8px;
        padding: 1rem;
        text-align: center;
        transition: all 0.3s ease;
    }
    .metric-item:hover {
        transform: translateY(-3px);
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.05);
    }
    .metric-label {
        color: #6b7280;
        font-size: 0.9rem;
        margin-bottom: 0.5rem;
    }
    .metric-value {
        font-size: 1.5rem;
        font-weight: 600;
        color: #1e3a8a;
    }
    
    /* Action Buttons */
    .action-buttons {
        display: flex;
        flex-direction: column;
        gap: 0.75rem;
    }
    .action-button {
        display: flex;
        align-items: center;
        padding: 0.75rem 1rem;
        border-radius: 8px;
        border: none;
        font-size: 1rem;
        font-weight: 500;
        cursor: pointer;
        transition: all 0.3s ease;
    }
    .action-button.primary {
        background-color: #4f46e5;
        color: white;
    }
    .action-button.secondary {
        background-color: #e0f2fe;
        color: #0369a1;
    }
    .action-button.tertiary {
        background-color: #f8fafc;
        color: #1e293b;
        border: 1px solid #e2e8f0;
    }
    .action-button:hover {
        transform: translateY(-2px);
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
    }
    .action-icon {
        margin-right: 0.75rem;
        font-size: 1.2rem;
    }
    
    /* Download options */
    .download-option {
        background-color: white;
        border-radius: 12px;
        padding: 1.5rem;
        height: 100%;
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.05);
        text-align: center;
        transition: all 0.3s ease;
    }
    .download-option:hover {
        transform: translateY(-5px);
        box-shadow: 0 10px 15px rgba(0, 0, 0, 0.1);
    }
    .download-icon {
        font-size: 2.5rem;
        margin-bottom: 1rem;
        color: #4f46e5;
    }
    .download-title {
        font-weight: 600;
        margin-bottom: 0.5rem;
        color: #1e3a8a;
    }
    .download-desc {
        color: #6b7280;
        font-size: 0.9rem;
        margin-bottom: 1rem;
    }
</style>

<!-- Add Lottie Player for animations -->
<script src="https://unpkg.com/@lottiefiles/lottie-player@latest/dist/lottie-player.js"></script>
""", unsafe_allow_html=True)

# Title and introduction for sidebar - ENHANCED MODERN UI
with st.sidebar:
    # Add custom CSS for the sidebar
    st.markdown("""
    <style>
    /* Sidebar styling */
    section[data-testid="stSidebar"] {
        background-color: #f8fafc;
        background-image: linear-gradient(135deg, #f8fafc 0%, #e0e7ff 100%);
        border-right: 0;
    }
    
    /* Sidebar title styling */
    .sidebar-title {
        color: #1e3a8a;
        font-weight: 700;
        font-size: 1.6rem;
        padding-bottom: 0.5rem;
        margin-bottom: 1.5rem;
        border-bottom: 2px solid #4f46e5;
        display: flex;
        align-items: center;
        gap: 10px;
    }
    
    /* Agent cards */
    .agent-sidebar-card {
        background-color: white;
        border-radius: 12px;
        padding: 1rem;
        margin-bottom: 1rem;
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.05);
        transition: all 0.3s ease;
        position: relative;
        overflow: hidden;
    }
    
    .agent-sidebar-card:hover {
        transform: translateY(-5px);
        box-shadow: 0 10px 15px rgba(0, 0, 0, 0.1);
    }
    
    .agent-sidebar-card::before {
        content: "";
        position: absolute;
        top: 0;
        left: 0;
        width: 5px;
        height: 100%;
        background: linear-gradient(to bottom, #4f46e5, #7c3aed);
        border-radius: 5px 0 0 5px;
    }
    
    .agent-icon {
        width: 40px;
        height: 40px;
        border-radius: 50%;
        background: linear-gradient(135deg, #4f46e5 0%, #7c3aed 100%);
        display: flex;
        align-items: center;
        justify-content: center;
        margin-bottom: 0.75rem;
        color: white;
        font-size: 1.25rem;
        position: relative;
        z-index: 1;
    }
    
    .agent-icon::after {
        content: "";
        position: absolute;
        width: 100%;
        height: 100%;
        background: linear-gradient(135deg, #4f46e5 0%, #7c3aed 100%);
        border-radius: 50%;
        opacity: 0.2;
        z-index: -1;
        animation: pulse 2s infinite;
    }
    
    @keyframes pulse {
        0% {
            transform: scale(1);
            opacity: 0.2;
        }
        50% {
            transform: scale(1.5);
            opacity: 0;
        }
        100% {
            transform: scale(1);
            opacity: 0.2;
        }
    }
    
    .agent-name {
        font-weight: 600;
        font-size: 1.1rem;
        color: #1e3a8a;
        margin-bottom: 0.25rem;
    }
    
    .agent-desc {
        font-size: 0.9rem;
        color: #6b7280;
        margin-bottom: 0.5rem;
    }
    
    /* Progress circles */
    .progress-circles {
        display: flex;
        gap: 5px;
    }
    
    .progress-circle {
        width: 8px;
        height: 8px;
        border-radius: 50%;
        background-color: #e5e7eb;
    }
    
    .progress-circle.active {
        background-color: #4f46e5;
    }
    
    /* Company brand card */
    .company-card {
        background-color: white;
        border-radius: 12px;
        padding: 1.25rem;
        margin-top: 1.5rem;
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.05);
        text-align: center;
    }
    
    .company-card img {
        margin-bottom: 0.75rem;
    }
    
    .copyright {
        font-size: 0.8rem;
        color: #6b7280;
        margin-top: 0.75rem;
    }
    
    /* Animation for sidebar */
    @keyframes fadeIn {
        from { opacity: 0; transform: translateX(-10px); }
        to { opacity: 1; transform: translateX(0); }
    }
    
    .agent-sidebar-card:nth-child(1) { animation: fadeIn 0.5s ease forwards; }
    .agent-sidebar-card:nth-child(2) { animation: fadeIn 0.5s ease forwards 0.1s; }
    .agent-sidebar-card:nth-child(3) { animation: fadeIn 0.5s ease forwards 0.2s; }
    .agent-sidebar-card:nth-child(4) { animation: fadeIn 0.5s ease forwards 0.3s; }
    
    /* Hide default Streamlit elements and spacer */
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    </style>
    """, unsafe_allow_html=True)
    
    # Modern sidebar header with icon
    st.markdown("""
    <div class="sidebar-title">
        <svg xmlns="http://www.w3.org/2000/svg" width="24" height="24" viewBox="0 0 24 24" fill="none" stroke="#4f46e5" stroke-width="2" stroke-linecap="round" stroke-linejoin="round">
            <rect x="3" y="3" width="18" height="18" rx="2" ry="2"></rect>
            <line x1="3" y1="9" x2="21" y2="9"></line>
            <line x1="9" y1="21" x2="9" y2="9"></line>
        </svg>
        Multi-Agent System
    </div>
    """, unsafe_allow_html=True)
    
    # Agent Cards with modern UI
    st.markdown("""
    <div class="agent-sidebar-card">
        <div class="agent-icon">🔍</div>
        <div class="agent-name">Document Parser</div>
        <div class="agent-desc">Analyzes RFP documents & extracts requirements</div>
        <div class="progress-circles">
            <div class="progress-circle active"></div>
            <div class="progress-circle active"></div>
            <div class="progress-circle active"></div>
        </div>
    </div>
    
    <div class="agent-sidebar-card">
        <div class="agent-icon">🧠</div>
        <div class="agent-name">Knowledge Retrieval</div>
        <div class="agent-desc">Connects your experience to RFP requirements</div>
        <div class="progress-circles">
            <div class="progress-circle active"></div>
            <div class="progress-circle active"></div>
            <div class="progress-circle"></div>
        </div>
    </div>
    
    <div class="agent-sidebar-card">
        <div class="agent-icon">✍️</div>
        <div class="agent-name">Response Generator</div>
        <div class="agent-desc">Creates tailored content for each section</div>
        <div class="progress-circles">
            <div class="progress-circle active"></div>
            <div class="progress-circle"></div>
            <div class="progress-circle"></div>
        </div>
    </div>
    
    <div class="agent-sidebar-card">
        <div class="agent-icon">🔍</div>
        <div class="agent-name">Quality Control</div>
        <div class="agent-desc">Reviews for completeness and compliance</div>
        <div class="progress-circles">
            <div class="progress-circle"></div>
            <div class="progress-circle"></div>
            <div class="progress-circle"></div>
        </div>
    </div>
    """, unsafe_allow_html=True)
    
    # Updated company branding section
    st.markdown("""
    <div class="company-card">
        <img src="https://via.placeholder.com/200x60.png?text=YourCompany" width="150">
        <div class="copyright">© 2025 Your Company Name</div>
    </div>
    """, unsafe_allow_html=True)

# Initialize session state for tracking progress and results
if 'rfp_text' not in st.session_state:
    st.session_state.rfp_text = None
if 'requirements' not in st.session_state:
    st.session_state.requirements = None
if 'knowledge' not in st.session_state:
    st.session_state.knowledge = None
if 'response_draft' not in st.session_state:
    st.session_state.response_draft = None
if 'review' not in st.session_state:
    st.session_state.review = None
if 'processing_complete' not in st.session_state:
    st.session_state.processing_complete = False

# Define uploaded_file at the global level before using it
uploaded_file = None

# Landing page with uploader
if st.session_state.rfp_text is None and not st.session_state.processing_complete:
    # Hero section with animated illustration
    st.markdown("""
    <div class="hero-container">
        <div class="hero-content">
            <h1>RFP Response Assistant</h1>
            <p class="hero-subtitle">Transform hours of manual work into minutes with AI</p>
            <div class="hero-stats">
                <div class="stat-item">
                    <span class="stat-number">60%</span>
                    <span class="stat-label">Time Saved</span>
                </div>
                <div class="stat-item">
                    <span class="stat-number">100%</span>
                    <span class="stat-label">Requirement Coverage</span>
                </div>
                <div class="stat-item">
                    <span class="stat-number">4×</span>
                    <span class="stat-label">Faster Response</span>
                </div>
            </div>
        </div>
        <div class="hero-image">
            <!-- Since we can't embed actual Lottie animations in Streamlit Cloud without using custom components,
                 we'll use an emoji as a placeholder -->
            <div style="font-size: 5rem; text-align: center;">📝✨</div>
        </div>
    </div>
    """, unsafe_allow_html=True)
    
    # Upload section
    st.markdown('<div class="upload-container">', unsafe_allow_html=True)
    st.subheader("📄 Upload Your RFP Document")
    uploaded_file = st.file_uploader("Upload an RFP document (PDF) to begin the automated response process", type=["pdf"])
    
    # Use a more visual way to show instructions with three simple steps
    st.markdown("""
    <div class="steps-container">
        <div class="step-item">
            <div class="step-number">1</div>
            <div class="step-icon">📄</div>
            <div class="step-title">Upload RFP</div>
            <div class="step-desc">Upload your RFP document in PDF format</div>
        </div>
        <div class="step-item">
            <div class="step-number">2</div>
            <div class="step-icon">🤖</div>
            <div class="step-title">AI Processing</div>
            <div class="step-desc">Our agents analyze and create a response</div>
        </div>
        <div class="step-item">
            <div class="step-number">3</div>
            <div class="step-icon">📝</div>
            <div class="step-title">Download</div>
            <div class="step-desc">Get your professional RFP response</div>
        </div>
    </div>
    """, unsafe_allow_html=True)
    
    # Add a "Learn more" expandable section instead of showing all the details
    with st.expander("Learn more about how it works"):
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("""
            ### 🔍 Document Parser Agent
            - Analyzes RFP documents
            - Extracts key requirements and structure
            - Identifies deadlines and evaluation criteria
            
            ### 🧠 Knowledge Retrieval Agent
            - Searches company knowledge base
            - Finds relevant past projects and experience
            - Identifies suitable team members and qualifications
            """)
        
        with col2:
            st.markdown("""
            ### ✍️ Response Generator Agent
            - Creates tailored content for each section
            - Aligns proposal with requirements
            - Formats response professionally
            
            ### 🔍 Quality Control Agent
            - Reviews for completeness and compliance
            - Identifies gaps and improvement areas
            - Highlights sections needing human expertise
            """)
    
    st.markdown('</div>', unsafe_allow_html=True)
else:
    # Add a new file uploader if process isn't complete yet
    if not st.session_state.processing_complete:
        uploaded_file = st.file_uploader("Upload an RFP document (PDF)", type=["pdf"])

# Extract text from uploaded document
if uploaded_file is not None:
    # Extract text only if we haven't done it yet
    if st.session_state.rfp_text is None:
        # Save the uploaded file to a temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
            tmp_file.write(uploaded_file.getvalue())
            tmp_path = tmp_file.name
        
        # Extract text from the PDF
        with st.spinner("Extracting text from PDF..."):
            pdf_reader = PdfReader(tmp_path)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text()
            
            os.unlink(tmp_path)  # Remove the temporary file
            
        st.session_state.rfp_text = text
        
        # Show confirmation
        st.success(f"Successfully extracted {len(text)} characters from {uploaded_file.name}")

# Process RFP button - when clicked, start the process
if st.session_state.rfp_text is not None and not st.session_state.processing_complete:
    if st.button("Start Multi-Agent Process", key="start_process"):
        # Display animated agent card for Document Parser
        st.markdown(f'''
        <div class="agent-card">
            <div class="agent-header">🔍 Document Parser Agent</div>
            <div class="agent-status">Analyzing RFP document and extracting key requirements...</div>
            <div class="animated-progress">
                <div class="animated-progress-bar"></div>
            </div>
        </div>
        ''', unsafe_allow_html=True)
        
        # Create a placeholder for the success message
        success_placeholder1 = st.empty()
        
        # Process document parser agent
        with st.spinner("Extracting structured requirements..."):
            # Create prompt for Document Parser Agent
            parser_prompt = f"""You are a Document Parser Agent specialized in analyzing RFP documents. 
            Extract the following information from this RFP:
            1. Key requirements and deliverables
            2. Compliance needs
            3. Deadlines
            4. Evaluation criteria
            5. Required sections for the response
            
            Format your response as JSON with these sections as keys.
            
            RFP content:
            {st.session_state.rfp_text[:15000]}  # Limit content to avoid token limits
            """
            
            # Call Anthropic API with system prompt
            parser_system_prompt = "You are a Document Parser Agent that extracts structured information from RFP documents."
            parser_response = call_anthropic_api(
                parser_prompt, 
                max_tokens=2000, 
                temperature=0,
                system_prompt=parser_system_prompt
            )
            
            # Try to extract JSON from the response
# REPLACE WITH THIS IMPROVED VERSION which handles more edge cases:
try:
    # First attempt to find JSON via standard method
    start_idx = parser_response.find('{')
    end_idx = parser_response.rfind('}') + 1
    
    if start_idx >= 0 and end_idx > start_idx:
        json_str = parser_response[start_idx:end_idx]
        
        try:
            # Try to parse the JSON string
            st.session_state.requirements = json.loads(json_str)
        except json.JSONDecodeError:
            # If standard JSON parsing fails, try a more flexible approach
            # This handles cases where the JSON might have formatting issues
            import re
            # Clean up the string - remove any problematic characters
            cleaned_json = re.sub(r'[\n\r\t]', '', json_str)
            # Try again with cleaned string
            try:
                st.session_state.requirements = json.loads(cleaned_json)
            except:
                # If still failing, create a structured dictionary manually
                # This is a fallback to ensure something is displayed
                requirements_text = parser_response.replace('\n', ' ').strip()
                st.session_state.requirements = {
                    "Key_Requirements_And_Deliverables": {
                        "Extracted_Requirements": requirements_text
                    },
                    "Compliance_Needs": {
                        "Compliance_Requirements": "Extraction failed - please review document manually"
                    },
                    "Deadlines": {
                        "Submission_Deadline": "Extraction failed - please review document manually"
                    },
                    "Evaluation_Criteria": {
                        "Criteria": "Extraction failed - please review document manually"
                    },
                    "Required_Sections_For_The_Response": [
                        "Executive Summary", 
                        "Technical Approach", 
                        "Experience", 
                        "Team Composition", 
                        "Project Plan"
                    ]
                }
    else:
        # If no JSON-like structure is found, create a fallback structure
        st.session_state.requirements = {
            "Key_Requirements_And_Deliverables": {
                "Main_Requirements": parser_response[:500] + "..."
            },
            "error": "Could not extract proper JSON format - showing raw text extract"
        }
except Exception as e:
    # Create a basic fallback structure in case of any error
    st.session_state.requirements = {
        "Key_Requirements_And_Deliverables": {
            "Error": "An error occurred during processing"
        },
        "error": f"Error details: {str(e)}",
        "raw_response": parser_response[:200] + "..."  # Truncated to avoid overly large error messages
    }
        
        # Show success indicator with SVG animation
        success_placeholder1.markdown(f'''
        <div class="success-indicator">
            <svg class="checkmark" xmlns="http://www.w3.org/2000/svg" viewBox="0 0 52 52">
                <circle class="checkmark-circle" cx="26" cy="26" r="25" fill="none"/>
                <path class="checkmark-check" fill="none" d="M14.1 27.2l7.1 7.2 16.7-16.8"/>
            </svg>
            <p>✅ Document Parser Agent completed successfully!</p>
        </div>
        ''', unsafe_allow_html=True)
        
        # Display animated agent card for Knowledge Retrieval
        st.markdown(f'''
        <div class="agent-card">
            <div class="agent-header">🧠 Knowledge Retrieval Agent</div>
            <div class="agent-status">Searching company knowledge base for relevant information...</div>
            <div class="animated-progress">
                <div class="animated-progress-bar"></div>
            </div>
        </div>
        ''', unsafe_allow_html=True)
        
        # Create a placeholder for the success message
        success_placeholder2 = st.empty()
        
        # Process knowledge retrieval agent
        with st.spinner("Retrieving relevant knowledge and past projects..."):
            # Create prompt for Knowledge Retrieval Agent
            knowledge_prompt = f"""Given these RFP requirements, provide relevant information that should be included in our response:
            
            {str(st.session_state.requirements)}
            
            Include:
            1. Suggested past projects that demonstrate relevant experience
            2. Key team members who should be mentioned
            3. Standard service descriptions that match the requirements
            4. Relevant compliance certifications and credentials
            
            Format as a structured list of recommendations.
            """
            
            # Call Anthropic API with system prompt
            knowledge_system_prompt = "You are a Knowledge Retrieval Agent for a professional services firm in Australia. You find relevant information from the company's knowledge base."
            st.session_state.knowledge = call_anthropic_api(
                knowledge_prompt, 
                max_tokens=2000, 
                temperature=0.2,
                system_prompt=knowledge_system_prompt
            )
        
        # Show success indicator with SVG animation
        success_placeholder2.markdown(f'''
        <div class="success-indicator">
            <svg class="checkmark" xmlns="http://www.w3.org/2000/svg" viewBox="0 0 52 52">
                <circle class="checkmark-circle" cx="26" cy="26" r="25" fill="none"/>
                <path class="checkmark-check" fill="none" d="M14.1 27.2l7.1 7.2 16.7-16.8"/>
            </svg>
            <p>✅ Knowledge Retrieval Agent completed successfully!</p>
        </div>
        ''', unsafe_allow_html=True)
        
        # Display animated agent card for Response Generator
        st.markdown(f'''
        <div class="agent-card">
            <div class="agent-header">✍️ Response Generator Agent</div>
            <div class="agent-status">Creating draft response sections based on requirements and knowledge...</div>
            <div class="animated-progress">
                <div class="animated-progress-bar"></div>
            </div>
        </div>
        ''', unsafe_allow_html=True)
        
        # Create a placeholder for the success message
        success_placeholder3 = st.empty()
        
        # Process response generator agent
        with st.spinner("Generating comprehensive response draft..."):
            # Create prompt for Response Generator Agent
            response_prompt = f"""Create draft responses for an RFP based on these requirements and available knowledge:
            
            RFP Requirements:
            {str(st.session_state.requirements)}
            
            Available Knowledge:
            {st.session_state.knowledge}
            
            Generate professional, compelling draft responses for key sections of the RFP.
            Format each section with a clear heading and concise, value-focused content.
            Use markdown formatting for better readability.
            """
            
            # Call Anthropic API with system prompt
            response_system_prompt = "You are a Response Generator Agent for an Australian professional services firm. You create professional RFP response content."
            st.session_state.response_draft = call_anthropic_api(
                response_prompt, 
                max_tokens=3500, 
                temperature=0.4,
                system_prompt=response_system_prompt
            )
        
        # Show success indicator with SVG animation
        success_placeholder3.markdown(f'''
        <div class="success-indicator">
            <svg class="checkmark" xmlns="http://www.w3.org/2000/svg" viewBox="0 0 52 52">
                <circle class="checkmark-circle" cx="26" cy="26" r="25" fill="none"/>
                <path class="checkmark-check" fill="none" d="M14.1 27.2l7.1 7.2 16.7-16.8"/>
            </svg>
            <p>✅ Response Generator Agent completed successfully!</p>
        </div>
        ''', unsafe_allow_html=True)
        
        # Display animated agent card for Quality Control
        st.markdown(f'''
        <div class="agent-card">
            <div class="agent-header">🔍 Quality Control Agent</div>
            <div class="agent-status">Reviewing generated response for completeness and compliance...</div>
            <div class="animated-progress">
                <div class="animated-progress-bar"></div>
            </div>
        </div>
        ''', unsafe_allow_html=True)
        
        # Create a placeholder for the success message
        success_placeholder4 = st.empty()
        
        # Process quality control agent
        with st.spinner("Performing comprehensive quality review..."):
            # Create prompt for Quality Control Agent
            review_prompt = f"""Review this draft RFP response against the requirements and provide feedback:
            
            Draft Response:
            {st.session_state.response_draft}
            
            RFP Requirements:
            {str(st.session_state.requirements)}
            
            Provide feedback on:
            1. Completeness - Are all requirements addressed?
            2. Compliance - Does it meet all compliance needs?
            3. Consistency - Is the response consistent throughout?
            4. Areas for improvement
            5. Sections requiring human expert review
            
            Format your response in markdown with clear sections.
            """
            
            # Call Anthropic API with system prompt
            review_system_prompt = "You are a Quality Control Agent that reviews RFP responses for completeness, compliance, and quality."
            st.session_state.review = call_anthropic_api(
                review_prompt, 
                max_tokens=2000, 
                temperature=0,
                system_prompt=review_system_prompt
            )
        
        # Show success indicator with SVG animation
        success_placeholder4.markdown(f'''
        <div class="success-indicator">
            <svg class="checkmark" xmlns="http://www.w3.org/2000/svg" viewBox="0 0 52 52">
                <circle class="checkmark-circle" cx="26" cy="26" r="25" fill="none"/>
                <path class="checkmark-check" fill="none" d="M14.1 27.2l7.1 7.2 16.7-16.8"/>
            </svg>
            <p>✅ Quality Control Agent completed successfully!</p>
        </div>
        ''', unsafe_allow_html=True)
        
        # Mark processing as complete
        st.session_state.processing_complete = True
        
        # Force page refresh to show results
        st.rerun()

# Display results if processing is complete
if st.session_state.processing_complete:
    st.markdown('<div class="success-box">', unsafe_allow_html=True)
    st.subheader("✅ RFP Processing Complete!")
    st.markdown("All agents have successfully completed their tasks. Review the outputs below and download the complete response when you're ready.")
    st.markdown("</div>", unsafe_allow_html=True)
    
    # Create tabs for different outputs
    tab1, tab2, tab3, tab4 = st.tabs(["📋 Requirements", "📚 Knowledge", "📝 Response Draft", "🔍 Quality Review"])
    
    with tab1:
        st.subheader("Extracted Requirements")
        if isinstance(st.session_state.requirements, dict):
            st.write("Requirements data structure is valid")
        else: 
            st.error("Requirements data is not in dictionary format")
            st.write(f"Type: {type(st.session_state.requirements)}")

    # Then make sure the display logic is robust with proper checks:
        if isinstance(st.session_state.requirements, dict):
    if "error" in st.session_state.requirements:
        st.error(st.session_state.requirements["error"])
        if "raw_response" in st.session_state.requirements:
            st.text(st.session_state.requirements["raw_response"])
    else:
        # Before trying to access specific keys, check if they exist
        req_keys = ["Key_Requirements_And_Deliverables", "Compliance_Needs", 
                   "Deadlines", "Evaluation_Criteria", "Required_Sections_For_The_Response"]
        
        for key in req_keys:
            if key in st.session_state.requirements:
                st.markdown(f'<div class="card">', unsafe_allow_html=True)
                
                # Use appropriate emoji and title based on key
                title_mapping = {
                    "Key_Requirements_And_Deliverables": "📋 Key Requirements and Deliverables",
                    "Compliance_Needs": "⚖️ Compliance Requirements",
                    "Deadlines": "⏱️ Critical Deadlines",
                    "Evaluation_Criteria": "🔍 Evaluation Criteria",
                    "Required_Sections_For_The_Response": "📑 Required Response Sections"
                }
                
                section_title = title_mapping.get(key, key.replace('_', ' ').title())
                st.markdown(f"### {section_title}")
                
                # Add appropriate subtitle
                subtitle_mapping = {
                    "Key_Requirements_And_Deliverables": "*The core requirements that must be addressed in your response:*",
                    "Compliance_Needs": "*Mandatory compliance aspects that must be addressed:*",
                    "Deadlines": "*Important dates and timeline requirements:*",
                    "Evaluation_Criteria": "*How your proposal will be evaluated:*",
                    "Required_Sections_For_The_Response": "*Sections that must be included in your proposal:*"
                }
                
                if key in subtitle_mapping:
                    st.markdown(subtitle_mapping[key])
                
                # Display data based on type
                data = st.session_state.requirements[key]
                
                # Handle different data structures appropriately
                if isinstance(data, dict):
                    for category, items in data.items():
                        category_name = category.replace('_', ' ').title()
                        st.markdown(f'<h4 class="json-key">{category_name}</h4>', unsafe_allow_html=True)
                        
                        if isinstance(items, list):
                            for item in items:
                                st.markdown(f'<div class="json-list-item">{item}</div>', unsafe_allow_html=True)
                        else:
                            st.markdown(f'<div class="json-value">{items}</div>', unsafe_allow_html=True)
                        st.markdown("---")
                elif isinstance(data, list):
                    # Create a table for required sections
                    sections_df = pd.DataFrame({
                        "Section Number": range(1, len(data) + 1),
                        "Section Name": data
                    })
                    st.dataframe(sections_df, use_container_width=True)
                else:
                    # If it's just a string or other type
                    st.markdown(str(data))
                
                st.markdown('</div>', unsafe_allow_html=True)
            # If a key is missing, create a minimal placeholder
            else:
                st.markdown(f'<div class="card">', unsafe_allow_html=True)
                section_title = key.replace('_', ' ').title()
                st.markdown(f"### {section_title}")
                st.markdown("*No data extracted for this section*")
                st.markdown('</div>', unsafe_allow_html=True)
else:
    # Fallback for completely invalid requirements
    st.markdown("No structured requirements could be extracted. Please try another document or check system configuration.")
    
    with tab2:
        st.subheader("Relevant Knowledge")
        st.markdown('<div class="card">', unsafe_allow_html=True)
        st.markdown(st.session_state.knowledge)
        st.markdown('</div>', unsafe_allow_html=True)
    
    with tab3:
        st.subheader("Generated Response Draft")
        st.markdown('<div class="card">', unsafe_allow_html=True)
        st.markdown(st.session_state.response_draft)
        st.markdown('</div>', unsafe_allow_html=True)
    
    with tab4:
        st.subheader("Quality Review")
        st.markdown('<div class="card">', unsafe_allow_html=True)
        st.markdown(st.session_state.review)
        st.markdown('</div>', unsafe_allow_html=True)
    
    # Modern dashboard header
    st.markdown("""
    <div class="dashboard-header">
        <h2>RFP Response Dashboard</h2>
        <p>Complete analysis and response for your RFP document</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Create metrics dashboard with visual charts
    col1, col2 = st.columns([2, 1])
    
    with col1:
        st.markdown('<div class="dashboard-card">', unsafe_allow_html=True)
        st.subheader("🔍 Response Analytics")
        
        # Count requirements
        req_count = 0
        if isinstance(st.session_state.requirements, dict) and "Key_Requirements_And_Deliverables" in st.session_state.requirements:
            req_data = st.session_state.requirements["Key_Requirements_And_Deliverables"]
            if isinstance(req_data, dict):
                for category, items in req_data.items():
                    if isinstance(items, list):
                        req_count += len(items)
                    else:
                        req_count += 1
        
        # Estimate response completeness
        completeness = "N/A"
        if st.session_state.review:
            # Look for percentages in the review text
            percentages = re.findall(r'(\d+)%', st.session_state.review)
            if percentages:
                try:
                    # Use the first percentage found as completeness
                    completeness = f"{percentages[0]}%"
                except:
                    pass
        
        # Count sections in response
        section_count = 0
        if st.session_state.response_draft:
            # Count markdown headings as sections
            section_count = st.session_state.response_draft.count('\n#')
        
        # Estimate time saved
        # Assume 30 minutes per requirement for manual processing
        time_saved = req_count * 30  # minutes
        if time_saved > 60:
            time_saved_str = f"{time_saved // 60}h {time_saved % 60}m"
        else:
            time_saved_str = f"{time_saved}m"
        
        # Create metrics in a more visual way
        metrics_data = {
            "Requirements": req_count,
            "Sections": section_count,
            "Completeness": completeness if isinstance(completeness, str) else f"{completeness}%",
            "Time Saved": time_saved_str
        }
        
        # Display as a horizontal metric bar
        st.markdown("""
        <div class="metrics-container">
        """, unsafe_allow_html=True)
        
        for label, value in metrics_data.items():
            st.markdown(f"""
            <div class="metric-item">
                <div class="metric-label">{label}</div>
                <div class="metric-value">{value}</div>
            </div>
            """, unsafe_allow_html=True)
        
        st.markdown("""
        </div>
        """, unsafe_allow_html=True)
        
        st.markdown('</div>', unsafe_allow_html=True)
    
    with col2:
        st.markdown('<div class="dashboard-card">', unsafe_allow_html=True)
        st.subheader("⚡ Quick Actions")
        
        # Display buttons for quick actions (Note: in Streamlit these are for display only)
        st.markdown("""
        <div class="action-buttons">
            <button class="action-button primary">
                <span class="action-icon">📝</span>
                <span class="action-text">Edit Response</span>
            </button>
            <button class="action-button secondary">
                <span class="action-icon">📤</span>
                <span class="action-text">Share Response</span>
            </button>
            <button class="action-button tertiary">
                <span class="action-icon">🔄</span>
                <span class="action-text">Start New RFP</span>
            </button>
        </div>
        """, unsafe_allow_html=True)
        
        st.markdown('</div>', unsafe_allow_html=True)
    
    # Create downloads section
    st.markdown("### Export Options")
    
    # Create columns for the two download options
    col1, col2 = st.columns(2)
    
    with col1:
        # Create downloadable Word document
        st.markdown('<div class="download-option">', unsafe_allow_html=True)
        st.markdown('<div class="download-icon">📝</div>', unsafe_allow_html=True)
        st.markdown('<div class="download-title">Word Document</div>', unsafe_allow_html=True)
        st.markdown('<div class="download-desc">Download as an editable Word document for further customization</div>', unsafe_allow_html=True)
        
        try:
            docx_buffer = generate_docx(st.session_state.response_draft, st.session_state.review)
            
            st.download_button(
                label="📝 Download as Word",
                data=docx_buffer,
                file_name="rfp_response.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        except Exception as e:
            st.error(f"Error generating Word document: {str(e)}")
        
        st.markdown('</div>', unsafe_allow_html=True)
    
    with col2:
        # Create downloadable PDF document
        st.markdown('<div class="download-option">', unsafe_allow_html=True)
        st.markdown('<div class="download-icon">📑</div>', unsafe_allow_html=True)
        st.markdown('<div class="download-title">PDF Document</div>', unsafe_allow_html=True)
        st.markdown('<div class="download-desc">Download as a professionally formatted PDF document</div>', unsafe_allow_html=True)
        
        try:
            pdf_buffer = generate_pdf(st.session_state.response_draft, st.session_state.review)
            
            st.download_button(
                label="📑 Download as PDF",
                data=pdf_buffer,
                file_name="rfp_response.pdf",
                mime="application/pdf"
            )
        except Exception as e:
            st.error(f"Error generating PDF: {str(e)}")
        
        st.markdown('</div>', unsafe_allow_html=True)
    
    # Additional options
    st.markdown("### Next Steps")
    st.info("You can now review and refine the generated response before submitting.")
    
    if st.button("Start New RFP", key="reset"):
        # Reset all session state
        st.session_state.rfp_text = None
        st.session_state.requirements = None
        st.session_state.knowledge = None
        st.session_state.response_draft = None
        st.session_state.review = None
        st.session_state.processing_complete = False
        
        # Force page refresh
        st.rerun()
